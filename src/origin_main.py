import streamlit as st
import os
import tempfile
import pymupdf
import docx2txt  # Extract text from DOCX
import openai
import requests
from bs4 import BeautifulSoup
from docx import Document
from duckduckgo_search import DDGS
import re
from dotenv import load_dotenv
#import jwt  # Install with `pip install PyJWT`
import requests

# Load environment variables
load_dotenv()
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
if not OPENAI_API_KEY:
    st.error("❌ OpenAI API Key is missing. Please set it in your environment variables.")
    st.stop()

#secret_key = os.getenv("SECRET_KEY")



# Authentication is disabled for local testing
# query_params = st.query_params
# if "token" in query_params:
#     token = "".join(query_params["token"])
# else:
#     token = None

# if not token:
#     st.warning("⚠️ No authentication token found. Please log in.")
#     st.stop()

# Validate the token
#try:
#    decoded_data = jwt.decode(token, secret_key, algorithms=["HS256"])
#    user_id = decoded_data.get("user_id")
#    user_email = decoded_data.get("email")
 #   st.session_state["user"] = {"id": user_id, "email": user_email}
#except jwt.ExpiredSignatureError:
 #   st.error("⏳ Session expired. Please log in again.")
#    st.stop()
##except jwt.InvalidTokenError:
#    st.error("❌ Invalid authentication. Please log in again.")
#    st.stop()

#st.set_page_config(layout="wide")
#st.write(f"✅ Logged in as {user_email}")

# =============================================================================
# OpenAI Client and Helper Functions
# =============================================================================

client = openai.Client(api_key=OPENAI_API_KEY)

# =============================================================================
# Layout: Three-Column Layout for Inputs, Editing, and Preview
# =============================================================================

left_col, mid_col, right_col = st.columns([1, 2, 2])

# -----------------------------------------------------------------------------
# Left Panel: Company Inputs & RFP Upload (Scrollable)
# -----------------------------------------------------------------------------
with left_col:
    with st.container():
        st.header("Company Info")
        company_name = st.text_input("Your Company Name", help="Full name of your company")
        services_offered = st.text_area("Services & Solutions", help="Describe what your company offers")
        industry = st.text_input("Industry", help="Industry your company operates in")
        competitive_advantage = st.text_area("Competitive Advantage", help="What makes your company stand out?")
        previous_use_cases = st.text_area("Previous Use Cases", help="List relevant case studies or past projects.")
        team_expertise = st.text_area("Team & Expertise", help="Describe key team members and their expertise.")
        mission_values = st.text_area("Company Mission & Values", help="Summarize your company’s mission and values.")

        # File Upload
        st.header("Upload RFP")
        uploaded_file = st.file_uploader("Upload RFP Document (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"])

# -----------------------------------------------------------------------------
# File Extraction Function
# -----------------------------------------------------------------------------
def extract_text_from_file(uploaded_file):
    if uploaded_file is not None:
        ext = uploaded_file.name.split(".")[-1].lower()
        if ext == "pdf":
            # Using PyMuPDF to read the PDF bytes
            pdf_doc = pymupdf.open(stream=uploaded_file.getvalue(), filetype="pdf")
            return "\n".join([page.get_text("text") for page in pdf_doc])
        elif ext == "docx":
            temp_file_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
            with open(temp_file_path, "wb") as f:
                f.write(uploaded_file.getvalue())
            text = docx2txt.process(temp_file_path)
            os.remove(temp_file_path)
            return text
        elif ext == "txt":
            return uploaded_file.getvalue().decode("utf-8")
    return None

# -----------------------------------------------------------------------------
# GPT-4o Summarization Function
# -----------------------------------------------------------------------------
def generate_summary(text, prompt):
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": prompt},
            {"role": "user", "content": text}
        ]
    )
    return response.choices[0].message.content

# -----------------------------------------------------------------------------
# External Search Functions for Company Information
# -----------------------------------------------------------------------------
def search_duckduckgo(company_name):
    """Fetch company info using DuckDuckGo search (Latest Version)."""
    with DDGS() as ddgs:
        results = list(ddgs.text(company_name + " company overview", max_results=5))
    
    if results:
        return "\n\n".join([res["body"] for res in results if "body" in res])
    return None

def search_linkedin(company_name):
    """Scrape LinkedIn public search for company info."""
    search_url = f"https://www.linkedin.com/search/results/companies/?keywords={company_name.replace(' ', '%20')}"
    headers = {"User-Agent": "Mozilla/5.0"}

    response = requests.get(search_url, headers=headers)
    if response.status_code == 200:
        soup = BeautifulSoup(response.text, "html.parser")
        snippets = [s.get_text() for s in soup.find_all("p")][:3]  # Extract some text
        return "\n".join(snippets)
    return None

def search_crunchbase(company_name):
    """Scrape Crunchbase public search for company info."""
    search_url = f"https://www.crunchbase.com/search/organizations/field/organizations/num_funding_rounds/{company_name.replace(' ', '-').lower()}"
    headers = {"User-Agent": "Mozilla/5.0"}

    response = requests.get(search_url, headers=headers)
    if response.status_code == 200:
        soup = BeautifulSoup(response.text, "html.parser")
        snippets = [s.get_text() for s in soup.find_all("p")][:3]  # Extract some text
        return "\n".join(snippets)
    return None

def generate_company_summary(text):
    """Use OpenAI GPT-4o to summarize company information."""
    if not text:
        return "No relevant company information found."
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "Summarize this company's information into a structured overview. Include an overview, services, and potential pain points."},
            {"role": "user", "content": text},
        ]
    )
    return response.choices[0].message.content

def search_company_info(client_company_name):
    """Search multiple sources for company information and generate a summary."""
    # Validate the company name before proceeding.
    if not client_company_name or not client_company_name.strip():
        return "No valid client company name was extracted from the RFP."
    
    sources = {
        "DuckDuckGo": search_duckduckgo(client_company_name),
        "LinkedIn": search_linkedin(client_company_name),
        "Crunchbase": search_crunchbase(client_company_name),
    }

    # Combine all non-empty results.
    combined_info = "\n\n".join(f"**{source}:** {info}" for source, info in sources.items() if info)
    return generate_company_summary(combined_info)

# -----------------------------------------------------------------------------
# Extract RFP Data & Process Once
# -----------------------------------------------------------------------------
if uploaded_file and "rfp_extracted" not in st.session_state:
    rfp_text = extract_text_from_file(uploaded_file)
    if rfp_text:
        # Summarize RFP details.
        st.session_state["rfp_summary"] = generate_summary(
            rfp_text, 
            "Summarize this RFP document. Extract key details like dates, requirements, contact information, and scope."
        )

        # Extract the required proposal sections.
        required_sections_raw = generate_summary(
            rfp_text,
            "Extract only the sections that the vendor is required to include in their proposal as specified in this RFP. "
            "Do not list sections from the RFP itself. Only include the vendor's response structure. "
            "Return a clean bullet-point list, formatted as:\n\n"
            "- Executive Summary\n- Technical Approach\n- Pricing Proposal\n- Implementation Plan\n- Compliance Statements\n\n"
            "Do not include explanations or any other text."
        )
        # Clean extracted sections.
        st.session_state["proposal_sections"] = [line.strip("- ").strip() for line in required_sections_raw.split("\n") if line.strip()]

        # Extract the requesting company name with a refined prompt.
        extracted_name = generate_summary(
            rfp_text,
            "Extract and return only the exact name of the company that is requesting this proposal. Do not include any additional text or explanations."
        )
        # Optional: Use regex to extract a clean company name (assuming it starts with a capital letter)
        match = re.search(r"([A-Z][\w&\s]+)", extracted_name)
        if match:
            client_company_name = match.group(1).strip()
        else:
            client_company_name = extracted_name.strip()
        
        st.session_state["client_company_name"] = client_company_name

        # Debug: Print the extracted client company name.
        st.write(f"**Extracted Client Company Name:** {st.session_state['client_company_name']}")

        # Search for company info using multiple sources.
        st.session_state["client_company_info"] = search_company_info(st.session_state["client_company_name"])

        st.session_state["rfp_extracted"] = True
        st.session_state["generated_sections"] = {}
        st.session_state["current_section"] = 0

# -----------------------------------------------------------------------------
# Middle Panel: Editing Section & RFP Info Tabs (Scrollable)
# -----------------------------------------------------------------------------
with mid_col:
    with st.container():
        st.header("Proposal Editor")
        if "rfp_extracted" in st.session_state:
            tabs = st.tabs(["Extracted RFP Details", "Requesting Company Info", "Edit Proposal"])

            with tabs[0]:  # Extracted RFP Details
                st.subheader("RFP Summary")
                st.markdown(f"**{st.session_state['rfp_summary']}**")
                st.subheader("Required Proposal Sections")
                st.markdown("\n".join([f"- {s}" for s in st.session_state["proposal_sections"]]))

            with tabs[1]:  # Requesting Company Info
                st.subheader("Requesting Company Information")
                if "client_company_name" in st.session_state:
                    st.markdown(f"**Company Name:** {st.session_state['client_company_name']}")
                if "client_company_info" in st.session_state:
                    st.markdown(f"**Company Overview:**\n{st.session_state['client_company_info']}")

            with tabs[2]:  # Proposal Editing
                if "proposal_sections" in st.session_state and "current_section" in st.session_state:
                    total_sections = len(st.session_state["proposal_sections"])
                    progress = (st.session_state["current_section"] + 1) / total_sections
                    st.progress(progress)
                    current_section = st.session_state["proposal_sections"][st.session_state["current_section"]]

                    st.subheader(f"Editing: {current_section}")

                    if current_section not in st.session_state["generated_sections"]:
                        section_content = generate_summary(
                            text=st.session_state["rfp_summary"], 
                            prompt=f"""
You are an expert proposal writer. Generate content for the section "{current_section}" in a professional tone.

Use the following company details to tailor the response:
- **Company Name:** {company_name}
- **Industry:** {industry}
- **Services & Solutions:** {services_offered}
- **Competitive Advantage:** {competitive_advantage}
- **Previous Use Cases:** {previous_use_cases}
- **Team & Expertise:** {team_expertise}
- **Mission & Values:** {mission_values}

Ensure the response aligns with the RFP details and focuses on how our company uniquely meets the requirements.
The content should be structured clearly and emphasize the company's strengths.
Focus on how our company’s services and solutions address the potential pain points mentioned in the RFP.
Keep the content within 200 words.
You may use bullet points to make the content easier to read, if needed.
Format the response in well-structured paragraphs without redundant information.
                            """
                        )
                        st.session_state["generated_sections"][current_section] = section_content

                    modified_content = st.text_area(f"Modify {current_section}", st.session_state["generated_sections"][current_section], height=500)

                    col1, col2 = st.columns([1, 1])
                    if col1.button("Previous", disabled=st.session_state["current_section"] == 0):
                        st.session_state["generated_sections"][current_section] = modified_content
                        st.session_state["current_section"] -= 1
                        st.experimental_rerun()
                    if col2.button("Next", disabled=st.session_state["current_section"] == total_sections - 1):
                        st.session_state["generated_sections"][current_section] = modified_content
                        st.session_state["current_section"] += 1
                        st.experimental_rerun()

# -----------------------------------------------------------------------------
# Right Panel: Final Proposal Preview & Download
# -----------------------------------------------------------------------------
with right_col:
    with st.container():
        st.header("Final Proposal")
        if "generated_sections" in st.session_state and st.session_state["generated_sections"]:
            proposal_doc = Document()
            proposal_doc.add_heading("Proposal Document", level=1)
            proposal_doc.add_paragraph(f"Prepared for: {st.session_state['client_company_name']}\n")

            proposal_preview_list = []  # Use list to manage clean formatting
            for section, content in st.session_state["generated_sections"].items():
                proposal_doc.add_heading(section, level=2)  # Add section title
                proposal_doc.add_paragraph(content)
                # Append section content for preview.
                proposal_preview_list.append(f"### {section}\n{content}")

            # Save proposal to a temporary file for download.
            temp_docx_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
            proposal_doc.save(temp_docx_path)

            with open(temp_docx_path, "rb") as f:
                st.download_button("Download Proposal", f, file_name="Generated_Proposal.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

            st.markdown("\n\n".join(proposal_preview_list), unsafe_allow_html=True)