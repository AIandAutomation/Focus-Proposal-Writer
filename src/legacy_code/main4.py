import streamlit as st
import os
import json
import tempfile
import openai
import requests
import re
from bs4 import BeautifulSoup
from docx import Document
import pymupdf
import docx2txt
import spacy
from duckduckgo_search import DDGS
from dotenv import load_dotenv
import re
from docx import Document
from docx.shared import Pt

# =============================================================================
# Persistence Helper Functions
# =============================================================================

STATE_FILE = "session_state.json"

def load_state():
    """Load session state from a JSON file, excluding keys for widgets that don’t allow value assignment."""
    widget_keys_to_skip = {"client_files", "your_files", "nav_go_back", "btn_nav_go_back", "nav_save_next", "btn_nav_save_next", "start_over"}
    if os.path.exists(STATE_FILE):
        try:
            with open(STATE_FILE, "r") as f:
                state = json.load(f)
            for key, value in state.items():
                if key in widget_keys_to_skip:
                    continue
                if key not in st.session_state:
                    st.session_state[key] = value
        except json.JSONDecodeError:
            st.warning("Session state file is empty or contains invalid JSON. Initializing with default values.")
        except Exception as e:
            st.error(f"Error loading session state: {e}")

def save_state():
    """Save the current session state to a JSON file, skipping non-serializable values."""
    try:
        safe_state = {}
        for k, v in st.session_state.items():
            try:
                json.dumps(v)
                safe_state[k] = v
            except TypeError:
                pass
        with open(STATE_FILE, "w") as f:
            json.dump(safe_state, f)
    except Exception as e:
        st.error(f"Error saving session state: {e}")

# Initialize proposal_settings if not present
if "proposal_settings" not in st.session_state:
    st.session_state["proposal_settings"] = {"type": "Government", "tone": "Formal"}

def reset_app():
    # Remove the persistent session state file if it exists
    if os.path.exists(STATE_FILE):
        os.remove(STATE_FILE)
    # Preserve only the essential keys
    keys_to_preserve = {"client_info", "your_info", "proposal_settings"}
    for key in list(st.session_state.keys()):
        if key not in keys_to_preserve:
            del st.session_state[key]
    # Reinitialize required keys
    if "proposal_settings" not in st.session_state:
        st.session_state["proposal_settings"] = {"type": "Government", "tone": "Formal"}
    st.session_state["current_step"] = 1  # return to step 1
    save_state()
    if hasattr(st, "experimental_rerun"):
        try:
            st.experimental_rerun()
        except Exception:
            pass


# =============================================================================
# 1. Load Environment Variables & Initialize OpenAI
# =============================================================================
load_dotenv()
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
if not OPENAI_API_KEY:
    st.error("❌ OpenAI API Key is missing. Please set it in your environment variables.")
    st.stop()
openai.api_key = OPENAI_API_KEY

st.set_page_config(layout="wide", page_title="FocusGPT: Proposal Writer")
load_state()


# =============================================================================
# 2. Sidebar Navigation & Reset Controls
# =============================================================================
st.sidebar.title("Proposal Navigation")
steps = [
    "1: Document Upload & Company Info",
    "2: Proposal Type & Structure",
    "3: Proposal Generation & Review",
    "4: Export & Handoff"
]

if st.sidebar.button("Start Over", key="start_over"):
    reset_app()

if "current_step" not in st.session_state:
    st.session_state["current_step"] = 1

selected_step_label = st.sidebar.radio("Go to Step", steps, index=st.session_state["current_step"] - 1)
selected_step = int(selected_step_label.split(":")[0].strip())
if selected_step != st.session_state["current_step"]:
    st.session_state["current_step"] = selected_step
    save_state()
    try:
        st.experimental_rerun()
    except AttributeError:
        pass

def next_step():
    if st.session_state["current_step"] < 5:
        st.session_state["current_step"] += 1
        save_state()

def prev_step():
    if st.session_state["current_step"] > 1:
        st.session_state["current_step"] -= 1
        save_state()
        try:
            st.experimental_rerun()
        except AttributeError:
            pass

# =============================================================================
# 3. Helper Functions (File Extraction, Summaries, Searches, etc.)
# =============================================================================
def extract_text_from_file(uploaded_file):
    if uploaded_file is not None:
        ext = uploaded_file.name.split(".")[-1].lower()
        if ext == "pdf":
            pdf_doc = pymupdf.open(stream=uploaded_file.getvalue(), filetype="pdf")
            return "\n".join([page.get_text("text") for page in pdf_doc])
        elif ext == "docx":
            temp_file_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
            with open(temp_file_path, "wb") as f:
                f.write(uploaded_file.getvalue())
            text = docx2txt.process(temp_file_path)
            os.remove(temp_file_path)
            return text
        elif ext == "txt":
            return uploaded_file.getvalue().decode("utf-8")
    return None

def generate_summary(text, prompt):
    response = openai.chat.completions.create(
        model="gpt-4o-mini",  # Change to "o1" if preferred
        messages=[
            {"role": "system", "content": prompt},
            {"role": "user", "content": text}
        ]
    )
    return response.choices[0].message.content

def search_duckduckgo(company_name):
    with DDGS() as ddgs:
        results = list(ddgs.text(company_name + " company overview", max_results=5))
    if results:
        return "\n\n".join([res["body"] for res in results if "body" in res])
    return None

def search_linkedin(company_name):
    search_url = f"https://www.linkedin.com/search/results/companies/?keywords={company_name.replace(' ', '%20')}"
    headers = {"User-Agent": "Mozilla/5.0"}
    response = requests.get(search_url, headers=headers)
    if response.status_code == 200:
        soup = BeautifulSoup(response.text, "html.parser")
        snippets = [s.get_text() for s in soup.find_all("p")][:3]
        return "\n".join(snippets)
    return None

def search_crunchbase(company_name):
    search_url = f"https://www.crunchbase.com/search/organizations/field/organizations/num_funding_rounds/{company_name.replace(' ', '-').lower()}"
    headers = {"User-Agent": "Mozilla/5.0"}
    response = requests.get(search_url, headers=headers)
    if response.status_code == 200:
        soup = BeautifulSoup(response.text, "html.parser")
        snippets = [s.get_text() for s in soup.find_all("p")][:3]
        return "\n".join(snippets)
    return None

def generate_company_summary(text):
    if not text:
        return "No relevant company information found."
    response = openai.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": "Summarize this company's information into a structured overview. Include an overview, services, and potential pain points."},
            {"role": "user", "content": text},
        ]
    )
    return response.choices[0].message.content

def search_company_info(client_company_name):
    if not client_company_name or not client_company_name.strip():
        return "No valid client company name was extracted from the RFP."
    sources = {
        "DuckDuckGo": search_duckduckgo(client_company_name),
        "LinkedIn": search_linkedin(client_company_name),
        "Crunchbase": search_crunchbase(client_company_name),
    }
    combined_info = "\n\n".join(
        f"**{source}:** {info}" for source, info in sources.items() if info
    )
    return generate_company_summary(combined_info)

def check_compliance_and_best_practices(section_text):
    return "No major compliance issues found. Consider adding more cost details."

# UPDATE: Load spaCy model for industry identification
nlp = spacy.load("en_core_web_sm")
def identify_industry(text):
    # A simple keyword-based approach using spaCy
    industry_keywords = {
        "Healthcare": ["hospital", "medical", "health", "clinic"],
        "Technology": ["software", "IT", "technology", "computer"],
        "Finance": ["bank", "finance", "investment", "money"],
        "Manufacturing": ["factory", "manufacture", "production", "industrial"],
        "Government": ["government", "public sector", "federal", "state", "local"]
    }
    doc = nlp(text.lower())
    scores = {industry: 0 for industry in industry_keywords}
    for token in doc:
        for industry, keywords in industry_keywords.items():
            if token.text in keywords:
                scores[industry] += 1
    return max(scores, key=scores.get) if max(scores.values()) > 0 else "Unknown"

def extract_needs(text):
    # A simple need extraction: split by period, take longer sentences as proxies for needs
    sentences = [sent.strip() for sent in text.split(".") if len(sent.split()) > 5]
    sentences.sort(key=lambda s: len(s), reverse=True)
    return sentences[:5]

def rank_need(need_text):
    revenue_terms = ["profit", "revenue", "market", "cost", "sales", "ROI"]
    government_terms = ["compliance", "regulation", "security", "audit", "government", "public"]
    revenue_score = sum(need_text.lower().count(term) for term in revenue_terms)
    government_score = sum(need_text.lower().count(term) for term in government_terms)
    return revenue_score, government_score

def get_ranked_needs():
    texts = []
    for key in ["client_RFP", "client_SOW", "User_company_RFP", "User_company_SOW"]:
        if key in st.session_state["sources"]:
            texts.append(st.session_state["sources"][key])
    combined_text = "\n".join(texts)
    if not combined_text:
        return ""
    needs = extract_needs(combined_text)
    ranked = sorted(needs, key=lambda n: sum(rank_need(n)), reverse=True)
    return "\n".join([f"- {n}" for n in ranked[:3]])

def get_relevant_sources(section_name):
    # UPDATE: Mapping keys now use singular form to match our aggregation
    mapping = {
        "Executive Summary": ["client_Company Info", "User_Company Info", "client_News", "User_News"],
        "Technical Approach": ["client_SOW", "client_RFP", "client_Case Study", "User_company_SOW", "User_company_RFP", "User_company_Case Study"],
        "Compliance Statements": ["client_RFP", "client_Company Info", "User_Company Info"],
        "Implementation Plan": ["client_SOW", "client_RFP", "User_company_SOW", "User_company_RFP"],
    }
    relevant_keys = mapping.get(section_name, [])
    texts = []
    for key in relevant_keys:
        if key in st.session_state["sources"]:
            texts.append(st.session_state["sources"][key])
    return "\n\n".join(texts)

def add_markdown_content(doc, md_text):
    lines = md_text.splitlines()
    for line in lines:
        # Check for markdown headings (e.g., ### Heading)
        header_match = re.match(r'^(#{1,6})\s+(.*)', line)
        if header_match:
            level = len(header_match.group(1))
            text = header_match.group(2)
            doc.add_heading(text, level=level)
        else:
            # Replace bold syntax (**text**) with plain text 
            line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
            para = doc.add_paragraph(line)
            if para.runs:  # Check if there's at least one run
                para.runs[0].font.size = Pt(11)

def segment_rfp_sections(text):
    """
    Splits the RFP text into sections based on common headings.
    Returns a dict with section titles as keys and section content as values.
    """
    # Define common RFP headings (case-insensitive)
    headings = r"(?:Requirements|Scope of Work|Evaluation Criteria|Deliverables|Project Overview)"
    # Use regex to split on headings while capturing them
    pattern = re.compile(rf"(?P<heading>{headings}).*?(?=\n[A-Z]|\Z)", re.IGNORECASE | re.DOTALL)
    sections = {}
    for match in pattern.finditer(text):
        heading = match.group("heading").strip()
        content = match.group(0).strip()
        sections[heading] = content
    return sections

def extract_requirements_with_spacy(text):
    """
    Uses spaCy to extract sentences containing key requirement terms.
    """
    nlp = spacy.load("en_core_web_sm")
    doc = nlp(text)
    # Define keywords that often indicate requirements
    requirement_keywords = {"must", "required", "shall", "deliverable", "mandatory", "criteria"}
    extracted = []
    for sent in doc.sents:
        # Convert sentence to lowercase and check for keywords
        if any(keyword in sent.text.lower() for keyword in requirement_keywords):
            extracted.append(sent.text.strip())
    return extracted

def chunk_text_by_char(text, max_chars=12000):
    """
    Splits the text into chunks with a maximum number of characters.
    This is a rough approximation to avoid exceeding the token limit.
    """
    paragraphs = text.split("\n")
    chunks = []
    current_chunk = ""
    for para in paragraphs:
        # +1 for the newline
        if len(current_chunk) + len(para) + 1 > max_chars:
            chunks.append(current_chunk)
            current_chunk = para
        else:
            current_chunk += "\n" + para
    if current_chunk:
        chunks.append(current_chunk)
    return chunks

@st.cache_data(show_spinner=False)
def hierarchical_generate_summary(text, prompt, max_chars=12000):
    # Split the text into smaller chunks
    chunks = chunk_text_by_char(text, max_chars)
    summaries = []
    
    # Generate summary for each chunk
    for chunk in chunks:
        summary = generate_summary(chunk, prompt)
        summaries.append(summary)
    
    # Combine the chunk summaries
    combined_summary_text = "\n".join(summaries)
    
    # Optionally, summarize the combined text if it's still too long
    final_summary = generate_summary(combined_summary_text, prompt)
    return final_summary


def summarize_debug_text(text, max_len=500):
    if len(text) > max_len:
        return text[:max_len] + "\n...[truncated]"
    return text

# =============================================================================
# 4. Main Workflow Steps (Two-Panel Layout)
# =============================================================================

# ----- Step 1: Document Upload & Company Info -----
if "sources" not in st.session_state:
    st.session_state["sources"] = {}

def section_1():
    st.title("Step 1: Document Upload & Company Info")
    col_left, col_right = st.columns(2)

    # Overall Navigation at Top Right
    nav_cols = st.columns([4, 1])
    with nav_cols[1]:
        if st.button("Save & Next"):
            st.success("Client and Company info auto-saved!")
            next_step()

    hide_file_uploader_style = """
    <style>
    .css-1wiv7s5 {
        display: none;
    }
    </style>
    """
    st.markdown(hide_file_uploader_style, unsafe_allow_html=True)
    
    col_left, col_right = st.columns(2)

    # LEFT PANEL: Client Company Info
    with col_left:
        st.header("Client Company Info")
        if "client_info" not in st.session_state:
            st.session_state["client_info"] = {"name": "", "description": "", "files": []}
        client = st.session_state["client_info"]
        client["name"] = st.text_input("Client Company Name", value=client.get("name", ""))
        client["description"] = st.text_area("Client Company Description", value=client.get("description", ""))
        
        uploaded_client_docs = st.file_uploader(
            "Upload Client Documents (e.g., RFP, SOW, Company Info, Case Study, News)",
            type=["pdf", "docx", "txt"],
            key="client_files",
            accept_multiple_files=True
        )
        
        client_files = []
        if uploaded_client_docs:
            st.markdown("### Uploaded Client Documents")
            for file in uploaded_client_docs:
                cols = st.columns([3, 3])
                with cols[0]:
                    st.write(file.name)
                with cols[1]:
                    # Use singular category names for consistency
                    category = st.selectbox(
                        "Category",
                        options=["RFP", "SOW", "Company Info", "Case Study", "News", "Whitepaper", "Other"],
                        key=f"client_category_{file.name}"
                    )
                text = extract_text_from_file(file)
                client_files.append({"name": file.name, "text": text, "category": category})
                key = "client_" + category  # aggregation key e.g., "client_RFP"
                if key in st.session_state["sources"]:
                    st.session_state["sources"][key] += "\n" + text
                else:
                    st.session_state["sources"][key] = text
            client["files"] = client_files
        st.session_state["client_info"] = client
    
    # RIGHT PANEL: Your Company Info
    with col_right:
        st.header("Your Company Info")
        if "your_info" not in st.session_state:
            st.session_state["your_info"] = {"name": "", "description": "", "files": []}
        your = st.session_state["your_info"]
        your["name"] = st.text_input("Your Company Name", value=your.get("name", ""))
        your["description"] = st.text_area("Your Company Description", value=your.get("description", ""))
        
        uploaded_your_docs = st.file_uploader(
            "Upload Your Company Documents (e.g., RFP, SOW, Company Info, Case Study, News, Whitepaper)",
            type=["pdf", "docx", "txt"],
            key="your_files",
            accept_multiple_files=True
        )
        
        your_files = []
        if uploaded_your_docs:
            st.markdown("### Uploaded Company Documents")
            for file in uploaded_your_docs:
                cols = st.columns([3, 3])
                with cols[0]:
                    st.write(file.name)
                with cols[1]:
                    category = st.selectbox(
                        "Category",
                        options=["RFP", "SOW", "Company Info", "Case Study", "News", "Whitepaper", "Other"],
                        key=f"your_category_{file.name}"
                    )
                text = extract_text_from_file(file)
                your_files.append({"name": file.name, "text": text, "category": category})
                key = "company_" + category  # using "company_" prefix
                if key in st.session_state["sources"]:
                    st.session_state["sources"][key] += "\n" + text
                else:
                    st.session_state["sources"][key] = text
            your["files"] = your_files
        st.session_state["your_info"] = your

# ----- Step 2: Proposal Type & Structure -----
def section_2():
    st.title("Step 2: Proposal Type & Structure")

    # Navigation Buttons at Top Right (unchanged)
    nav_cols = st.columns(2)
    with nav_cols[0]:
        if st.button("Go Back", key="btn_nav_go_back"):
            prev_step()
    with nav_cols[1]:
        if st.button("Save & Next", key="btn_nav_save_next"):
            next_step()

    # Inject custom CSS for navigation buttons (unchanged)
    st.markdown(
        """
        <style>
        div[data-testid="stHorizontalBlock"] > div:nth-child(1) > button {
            background-color: #FF4136 !important;
            color: white !important;
            border: none !important;
            padding: 8px 16px !important;
            border-radius: 4px;
        }
        div[data-testid="stHorizontalBlock"] > div:nth-child(2) > button {
            background-color: #2ECC40 !important;
            color: white !important;
            border: none !important;
            padding: 8px 16px !important;
            border-radius: 4px;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )
    
    col_left, col_right = st.columns(2)

    # ---------------------------
    # LEFT PANEL: Set Proposal Type & Tone
    # ---------------------------
    with col_left:
        st.header("Set Proposal Type")
        if "proposal_settings" not in st.session_state:
            st.session_state["proposal_settings"] = {"type": "Government", "tone": "Formal"}
        settings = st.session_state["proposal_settings"]
        settings["type"] = st.radio("Proposal Type",
                                    ["Government", "Commercial", "Custom"],
                                    index=["Government", "Commercial", "Custom"].index(settings.get("type", "Government")),
                                    key="sec2_proposal_type")
        settings["tone"] = st.selectbox("Proposal Tone",
                                        ["Formal", "Persuasive", "Concise", "Custom"],
                                        index=["Formal", "Persuasive", "Concise", "Custom"].index(settings.get("tone", "Formal")),
                                        key="sec2_proposal_tone")
                                        
        # ---------------------------
        # NEW: Flexible Extraction from Client RFP/SOW Documents
        # ---------------------------
        # First, aggregate texts from client_RFP and client_SOW keys.
        aggregated_rfp_sow = ""
        for key in ["client_RFP", "client_SOW"]:
            if key in st.session_state["sources"]:
                aggregated_rfp_sow += "\n" + st.session_state["sources"][key]

        # Debug: show aggregated content and available source keys
        debug_full = st.checkbox("Show full aggregated content", value=False)

        if debug_full:
            st.write("DEBUG: Aggregated RFP/SOW Content:", aggregated_rfp_sow)
        else:
            st.write("DEBUG: Aggregated RFP/SOW Content:", summarize_debug_text(aggregated_rfp_sow))
        st.write("DEBUG: Available Source Keys:", list(st.session_state["sources"].keys()))

        # --- Document Type Detection & Extraction Configuration ---
        if aggregated_rfp_sow:
            def determine_document_type(text):
                text_upper = text.upper()
                if text_upper.count("RFP") > text_upper.count("RFQ"):
                    return "RFP"
                elif "RFQ" in text_upper:
                    return "RFQ"
                else:
                    return "General"
            
            # Define extraction templates for different document types.
            extraction_templates = {
                "RFP": {
                    "prompt": """
You are an expert proposal analyst. Analyze the provided RFP document text and extract a structured, bullet-point list of:
• Key requirements
• Required documents (if mentioned)
Format your output so that each item starts with "- " (a dash and a space) and include no additional text.
""",
                },
                "RFQ": {
                    "prompt": """
You are an expert procurement analyst. Analyze the provided RFQ document text and extract a structured, bullet-point list that includes:
• Key technical requirements
• Key commercial requirements
Format your output as bullet points with "- " at the start of each line.
""",
                },
                "General": {
                    "prompt": """
Analyze the provided document text and extract a concise bullet-point list of the most important requirements and components.
Ensure that each item begins with "- " (a dash followed by a space) and strictly list only the items.
""",
                }
            }
            
            detected_type = determine_document_type(aggregated_rfp_sow)
            # Allow the user to confirm or adjust the document type
            doc_type = st.selectbox("Select Document Type", options=["RFP", "RFQ", "General"], 
                                      index=["RFP", "RFQ", "General"].index(detected_type))
            extraction_prompt = extraction_templates[doc_type]["prompt"]
            
            # Use hierarchical summarization (with caching) to process large input text.
            extracted_req = hierarchical_generate_summary(aggregated_rfp_sow, extraction_prompt)
            # Convert the output into a list of requirements.
            extracted_list = [line.strip("- ").strip() for line in extracted_req.splitlines() if line.strip()]
            st.session_state["extracted_requirements"] = extracted_list
        else:
            st.write("DEBUG: No aggregated RFP/SOW content found.")

        # --- Additional Extraction from Client RFP Only ---
        aggregated_rfp = st.session_state["sources"].get("client_RFP", "")
        if aggregated_rfp:
            # Structured Segmentation: Isolate the "Requirements" section.
            segmented_sections = segment_rfp_sections(aggregated_rfp)
            requirements_section = ""
            for heading, content in segmented_sections.items():
                if "requirement" in heading.lower():
                    requirements_section = content
                    break
            # Fallback to the entire text if no specific section is found.
            if not requirements_section:
                requirements_section = aggregated_rfp

            req_prompt_secondary = """
You are an expert proposal analyst. Extract a concise bullet-point list of key requirements from the following RFP section. Only list the items.
            """
            # Use hierarchical summarization if the section is very long.
            if len(requirements_section) > 10000:
                extracted_req_secondary = hierarchical_generate_summary(requirements_section, req_prompt_secondary)
            else:
                extracted_req_secondary = generate_summary(requirements_section, req_prompt_secondary)
            extracted_list_secondary = [line.strip("- ").strip() for line in extracted_req_secondary.splitlines() if line.strip()]

            # Use spaCy-based extraction for additional detail.
            extracted_list_spacy = extract_requirements_with_spacy(requirements_section)

            # Combine and deduplicate both sets of results.
            combined_requirements = list(set(extracted_list_secondary + extracted_list_spacy))
            st.session_state["extracted_requirements"] = combined_requirements

        # User Feedback: Display the extracted requirements for review and adjustment.
        st.markdown("#### Extracted Requirements:")
        default_reqs = "\n".join(f"- {req}" for req in st.session_state.get("extracted_requirements", []))
        user_adjusted_reqs = st.text_area("Review and adjust the extracted requirements:", value=default_reqs, height=150)
        final_reqs = [line.strip("- ").strip() for line in user_adjusted_reqs.splitlines() if line.strip()]
        st.session_state["extracted_requirements"] = final_reqs

    # ---------------------------
    # RIGHT PANEL: Proposal Structure + Model's Understanding of Client Company
    # ---------------------------
    with col_right:
        st.header("Proposal Structure")
        default_outline = "\n\n- Executive Summary\n- Technical Approach\n- Pricing Proposal\n- Implementation Plan\n- Compliance Statements\n- Conclusion\n\n"
        outline = st.text_area("Edit or add sections in your proposal outline:",
                               value=st.session_state.get("proposal_sections", default_outline),
                               height=200,
                               key="sec2_outline")
        st.session_state["proposal_sections"] = outline
        
        # NEW: Load model's understanding of client company info
        client_description = st.session_state["client_info"].get("description", "")
        client_docs = st.session_state["sources"].get("client_Company Info", "")
        combined_text = client_description + "\n\n" + client_docs

        if st.button("Analyze Client Organization", key="analyze_client"):
            if combined_text.strip():
                company_analysis = generate_company_summary(combined_text)
                st.markdown("#### Client Organization Analysis:")
                st.write(company_analysis)
            else:
                st.write("Insufficient data to analyze organization info.")


# ---------------------------
# (Existing global variable computation for Step 3 is assumed to follow here)
# ---------------------------
sections_list = [sec.lstrip("- ").strip() for sec in st.session_state.get("proposal_sections", "").splitlines() if sec.strip()]
current_section = sections_list[st.session_state.get("current_section", 0)] if sections_list else "Undefined Section"
aggregated_sources = get_relevant_sources(current_section)
needs_summary = get_ranked_needs()

company_name = st.session_state.get("your_info", {}).get("name", "").strip() or "Unknown Company"
industry = st.session_state.get("industry", "").strip() or st.session_state.get("identified_industry", "Unknown")
services = st.session_state.get("services_offered", "").strip() or "N/A"
advantage = st.session_state.get("competitive_advantage", "").strip() or "N/A"
use_cases = st.session_state.get("previous_use_cases", "").strip() or "N/A"
team = st.session_state.get("team_expertise", "").strip() or "N/A"
mission = st.session_state.get("mission_values", "").strip() or "N/A"
tone = st.session_state["proposal_settings"].get("tone", "Formal")

# ----- Step 3: Proposal Generation & Review -----
def section_3():
    st.title("Step 3: Proposal Generation & Review")
    nav_cols = st.columns(2)
    with nav_cols[0]:
        if st.button("Go Back", key="btn_nav_go_back_sec3"):
            prev_step()
    with nav_cols[1]:
        if st.button("Save & Next", key="btn_nav_save_next_sec3"):
            next_step()
    col_left, col_right = st.columns(2)
    with col_left:
        st.header("Draft Proposal Content")
        section_nav = st.columns(2)
        with section_nav[0]:
            if st.button("Previous Section"):
                if st.session_state.get("current_section", 0) > 0:
                    st.session_state["current_section"] -= 1
                    if hasattr(st, "experimental_rerun"):
                        st.experimental_rerun()
                    else:
                        st.info("Please refresh the page to see updates.")
        with section_nav[1]:
            outline_text = st.session_state.get("proposal_sections", "")
            sections_list = [sec.lstrip("- ").strip() for sec in outline_text.splitlines() if sec.strip()]
            if st.button("Next Section"):
                if st.session_state.get("current_section", 0) < len(sections_list) - 1:
                    st.session_state["current_section"] += 1
                    if hasattr(st, "experimental_rerun"):
                        st.experimental_rerun()
                    else:
                        st.info("Please refresh the page to see updates.")
        outline_text = st.session_state.get("proposal_sections", "")
        sections_list = [sec.lstrip("- ").strip() for sec in outline_text.splitlines() if sec.strip()]
        current_index = st.session_state.get("current_section", 0)
        if current_index >= len(sections_list):
            current_index = 0
        selected_section = st.radio("Select a section to edit:", sections_list, index=current_index)
        st.session_state["current_section"] = sections_list.index(selected_section)
        current_section = selected_section
        st.write("Editing: " + current_section)
        if "generated_sections" not in st.session_state:
            st.session_state["generated_sections"] = {}
        base_prompt = f"""
You are an expert proposal writer. Generate content for the section "{current_section}" in a {tone} tone.
Use the following company details:
- Company Name: {company_name}
- Industry: {industry}
- Services & Solutions: {services}
- Competitive Advantage: {advantage}
- Previous Use Cases: {use_cases}
- Team & Expertise: {team}
- Mission & Values: {mission}

Below is the aggregated information from uploaded documents relevant to this section:
{aggregated_sources}

The following key needs have been identified based on the documents:
{needs_summary}
"""
        deck_content_snippet = ""
        if st.session_state.get("company_deck_text"):
            deck_content_snippet = f"\nAdditional Deck/Resource Info:\n{st.session_state['company_deck_text']}\n"
        global_instructions = f"""
{deck_content_snippet}
Ensure the response aligns with the client company uploaded RFP file details, and focuses on how our company uniquely meets the requirements.
Keep the content within 500 words.
You may use bullet points if needed, and ensure the text is well-structured and not redundant.
"""
        if "Technical Approach" in current_section:
            # SECTION CHANGES START:
            aggregated_rfp = st.session_state.get("sources", {}).get("client_RFP", "")
            extracted_reqs = st.session_state.get("extracted_requirements", [])
            company_expertise = st.session_state.get("company_expertise", "Your company’s relevant experience details...")
            additional_guidance = f"""
You are an expert proposal analyst and solution architect. You have been provided with:

1. Key Requirements extracted from the RFP:
{chr(10).join(f"- {req}" for req in extracted_reqs)}

2. Original RFP Summary (for additional context):
\"\"\"{aggregated_rfp}\"\"\"

3. Company Expertise/Case Studies:
\"\"\"{company_expertise}\"\"\"

Your task is to:
1. Confirm & Expand on Requirements:
   - Summarize and clarify these requirements in your own words.
   - Identify any additional context or constraints mentioned implicitly.
2. Reason About Tech Stacks:
   - Propose suitable technologies, frameworks, and infrastructure components.
   - Explain why each recommended technology fits, noting performance, scalability, and security.
3. Leverage Company Expertise:
   - Incorporate relevant case studies and experience details.
4. Describe a Cohesive Solution Architecture:
   - Explain how the chosen tech stacks integrate to address each requirement.
   - Outline key milestones or phases if applicable.
5. Ensure Compliance:
   - Address any mandatory documents or compliance items from the requirements.
6. Identify Risks & Mitigation Strategies:
   - Highlight potential risks and their mitigations.
7. Conclude with Competitive Advantage:
   - Summarize why this approach is optimal.

Constraints:
- Write in a professional, clear, and logically structured manner.
- Only use the provided RFP summary, requirements, and company expertise for details.

Now, produce a comprehensive technical proposal section.
"""

        elif "Implementation Plan" in current_section or "Work Plan" in current_section:
            additional_guidance = """
For the Implementation Plan (or Work Plan & Schedule):
1. Provide a clear timeline or Gantt-like structure, referencing tasks, durations, and milestones.
2. Discuss acceptance criteria, deliverable sign-offs, or how you'll address SOW deliverables.
3. Outline how you'll manage resources, risks, and communications throughout the phases.
4. Reference any compliance or regulatory steps (like security audits or quality gates).
"""
        elif current_section == "Pricing Proposal":
            additional_guidance = """
For Pricing Proposal:
Based on the following RFP content and identified requirements, generate a detailed labor cost estimate for the proposal.
If the RFP specifies roles, responsibilities, or job titles, extract and use them; otherwise, construct a recommended team that meets the RFP's needs.
Present: 
1. techstacks needed for each position to meet the feature, and reason number of hours needed.
2. the cost estimate in a table format (plain text) with these columns:
- Role
- Hourly Rate ($)
- Estimated Hours
- Total Cost ($)  (Hourly Rate * Estimated Hours)
Include a final line with the total deliverable cost.
For example:

Role                  Hourly Rate     Estimated Hours     Total Cost
Sr. Web Developer     116             1600                185600
Jr. Web Developer     85              600                 51000
...
Total Deliverable Cost: 236600
"""
        elif "Compliance" in current_section:
            additional_guidance = """
For Compliance Statements:
1. Mention relevant compliance frameworks (FISMA, NIST, HIPAA, etc.) if applicable to the RFP.
2. Address how you will complete any mandatory government forms (e.g., STD. 204, GSPD-05-105).
3. Discuss non-discrimination policies, meeting deadlines, and accountability for deliverables.
4. Show how your approach to data security and confidentiality aligns with SOW references.
"""
        elif "Maintenance" in current_section or "Support" in current_section:
            additional_guidance = """
For Maintenance & Support:
1. Provide details on post-launch support, help desk/ticketing, and escalation procedures.
2. Include a Service Level Agreement (SLA) structure (Critical, High, Medium, Low) with response times.
3. Mention patching, updates, ongoing training, and how you’ll keep the system secure and compliant.
4. Reference how you handle future enhancements or major changes in scope.
"""
        else:
            additional_guidance = """
Additionally, if relevant, please reference any SOW items or compliance standards that might apply.
Include best practices or structured details to make the proposal more aligned with typical Government RFP expectations.
"""
        final_prompt = base_prompt + global_instructions + additional_guidance
        if current_section not in st.session_state["generated_sections"]:
            section_content = generate_summary(
                text=st.session_state.get("rfp_summary", ""),
                prompt=final_prompt
            )
            st.session_state["generated_sections"][current_section] = section_content
        edited_draft = st.text_area("Edit Section Content:",
                                    value=st.session_state["generated_sections"][current_section],
                                    height=300)
        st.session_state["generated_sections"][current_section] = edited_draft
        st.write("FocusGPT Expert Suggestions:")
        if st.button("Generate FocusGPT Expert Suggestions"):
            focusgpt_prompt = f"""
You are FocusGPT, an expert proposal writer. Provide 3 to 4 concise bullet points through the client company reader's perspective that offer expert suggestions to improve the following section content.
Section Content:
{edited_draft}
"""
            focus_suggestions_text = generate_summary("", focusgpt_prompt)
            # Split bullet points assuming each starts with "-"
            suggestions = [line.strip("- ").strip() for line in focus_suggestions_text.splitlines() if line.strip()]
            st.session_state["focus_suggestions"] = suggestions

        if "focus_suggestions" in st.session_state:
            st.markdown("#### Select FocusGPT Suggestions to Apply:")
            selected_suggestions = []
            for i, suggestion in enumerate(st.session_state["focus_suggestions"]):
                if st.checkbox(suggestion, key=f"focus_suggestion_{i}", value=False):
                    selected_suggestions.append(suggestion)

            if st.button("Apply FocusGPT Suggestions"):
                if selected_suggestions:
                    suggestions_text = "\n".join(f"- {s}" for s in selected_suggestions)
                else:
                    suggestions_text = ""
                apply_prompt = f"""
You are FocusGPT, an expert proposal writer. Revise the following section content by incorporating the expert suggestions provided below. Produce a revised version that improves content clarity, structure, and compliance.
Expert Suggestions:
{suggestions_text}
Current Section Content:
{edited_draft}
"""
                revised_content = generate_summary("", apply_prompt)
                st.session_state["generated_sections"][current_section] = revised_content
                if hasattr(st, "experimental_rerun"):
                    st.experimental_rerun()
                else:
                    st.info("Please refresh the page to see updates.")
        st.write("User Feedback:")
        user_feedback = st.text_area("Enter your feedback to improve the section content:", value="", height=100, key="user_feedback")
        if st.button("Apply User Feedback"):
            feedback_prompt = f"""
You are an expert proposal writer. Revise the following section content based on the user feedback provided below. Produce a revised version that reflects the improvements requested.
User Feedback:
{user_feedback}
Current Section Content:
{edited_draft}
"""
            new_content = generate_summary("", feedback_prompt)
            st.session_state["generated_sections"][current_section] = new_content
            if hasattr(st, "experimental_rerun"):
                st.experimental_rerun()
            else:
                st.info("Please refresh the page to see updates.")
    with col_right:
        st.header("Live Proposal Preview")
        if st.button("Refresh Preview"):
            if hasattr(st, "experimental_rerun"):
                st.experimental_rerun()
            else:
                st.info("Please refresh the page to see updates.")
        preview_container = st.empty()
        preview_content = ""
        for sec, content in st.session_state.get("generated_sections", {}).items():
            preview_content += f"Section: {sec}\n{content}\n\n"
        preview_html = f'<div style="height:800px; overflow:auto; border:1px solid #ddd; padding:10px; white-space: pre-wrap;">{preview_content}</div>'
        st.markdown(preview_html, unsafe_allow_html=True)

# ----- Step 4: Export & Handoff -----
def section_4():
    st.title("Step 4: Export & Handoff")
    nav_cols = st.columns([4, 1])
    with nav_cols[1]:
        if st.button("Go Back"):
            prev_step()
    col_left, col_right = st.columns(2)
    
    # ----- Left Panel: Export Settings & Download -----
    with col_left:
        st.header("Export Settings & Final Touches")
        cover_page = st.text_input("Cover Page Text (optional):", value=st.session_state.get("cover_page", ""))
        st.session_state["cover_page"] = cover_page
        metadata = st.text_area("Enter metadata (submission date, contact info, version):",
                                value=st.session_state.get("metadata", ""))
        st.session_state["metadata"] = metadata

        # Generate the final proposal document for download
        final_doc = Document()
        #final_doc.add_heading("Final Proposal Document", level=1)
        client = st.session_state.get("client_info", {})
        your = st.session_state.get("your_info", {})
        final_doc.add_paragraph(f"Prepared for: {client.get('name', 'Client')}")
        final_doc.add_paragraph(f"Prepared by: {your.get('name', 'Your Company')}")
        if st.session_state.get("cover_page"):
            final_doc.add_paragraph(st.session_state["cover_page"])
        if st.session_state.get("metadata"):
            final_doc.add_paragraph(st.session_state["metadata"])
        for sec, content in st.session_state.get("generated_sections", {}).items():
            final_doc.add_heading(sec, level=2)
            add_markdown_content(final_doc, content)
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        final_doc.save(temp_file.name)
        st.download_button(
            "Download Final Proposal (DOCX)",
            open(temp_file.name, "rb"),
            file_name="Final_Proposal.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    
    # ----- Right Panel: Required Documents Checklist -----
    with col_right:
        st.header("Required Documents Checklist")
        if "extracted_requirements" in st.session_state:
            for req in st.session_state["extracted_requirements"]:
                st.checkbox(req, value=False, key="check_" + req)

# =============================================================================
# 5. Main App Dispatcher
# =============================================================================
def main():
    st.title("FocusGPT: Proposal Writer")
    current = st.session_state["current_step"]
    if current == 1:
        section_1()
    elif current == 2:
        section_2()
    elif current == 3:
        section_3()
    elif current == 4:
        section_4()
    save_state()

if __name__ == "__main__":
    main()
