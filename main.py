"""
Multi-Agent AI Proposal Drafting System (main.py)
-------------------------------------------
This application uses advanced AI agents to help users create compelling 
proposal documents for enterprise and government clients.
Run this app locally with:
    streamlit run main.py
"""

import os
import json
import tempfile
import streamlit as st
import ast
import re
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt
from datetime import datetime

from src.coordinator import CoordinatorAgent
from src.agents.document_extraction import DocumentExtractionAgent

# =============================================================================
# Persistence Helper Functions
# =============================================================================

STATE_FILE = "session_state.json"

def load_state():
    """Load session state from a JSON file, excluding keys for widgets that don't allow value assignment."""
    widget_keys_to_skip = {
        # File upload widgets
        "client_files", "your_files", 
        
        # Navigation buttons
        "nav_go_back", "btn_nav_go_back", "nav_save_next", "btn_nav_save_next", "start_over",
        "nav_back_section2", "nav_next_section2", "nav_back_section3", "nav_next_section3", "nav_back_section4",
        
        # Tab navigation buttons
        "continue_to_analysis", "proceed_to_brainstorm", "back_to_upload", "skip_to_next",
        
        # Section buttons
        "prev_section_btn", "next_section_btn", "refresh_preview_btn",
        
        # Action buttons
        "extract_key_info", "extract_req_btn", "gen_tech_btn", "gen_doc_btn", "download_btn"
    }
    
    if os.path.exists(STATE_FILE):
        try:
            with open(STATE_FILE, "r") as f:
                state = json.load(f)
            for key, value in state.items():
                if key in widget_keys_to_skip:
                    continue
                if key not in st.session_state:
                    st.session_state[key] = value
        except Exception as e:
            st.error(f"Error loading session state: {e}")

def save_state():
    """Save the current session state to a JSON file, skipping non-serializable values and widget keys."""
    try:
        # Use the same widget_keys_to_skip set as in load_state
        widget_keys_to_skip = {
            # File upload widgets
            "client_files", "your_files", 
            
            # Navigation buttons
            "nav_go_back", "btn_nav_go_back", "nav_save_next", "btn_nav_save_next", "start_over",
            "nav_back_section2", "nav_next_section2", "nav_back_section3", "nav_next_section3", "nav_back_section4",
            
            # Tab navigation buttons
            "continue_to_analysis", "proceed_to_brainstorm", "back_to_upload", "skip_to_next",
            
            # Section buttons
            "prev_section_btn", "next_section_btn", "refresh_preview_btn",
            
            # Action buttons
            "extract_key_info", "extract_req_btn", "gen_tech_btn", "gen_doc_btn", "download_btn"
        }
        
        safe_state = {}
        for k, v in st.session_state.items():
            # Skip widget keys
            if k in widget_keys_to_skip:
                continue
                
            # Also skip keys that start with dynamically generated patterns
            if (k.startswith("gen_btn_") or 
                k.startswith("clear_btn_") or 
                k.startswith("feedback_") or 
                k.startswith("edit_") or
                k.startswith("client_category_") or
                k.startswith("your_category_") or
                k.startswith("check_")):
                continue
                
            # Skip non-serializable values
            try:
                json.dumps(v)
                safe_state[k] = v
            except TypeError:
                pass
                
        with open(STATE_FILE, "w") as f:
            json.dump(safe_state, f)
    except Exception as e:
        st.error(f"Error saving session state: {e}")

def reset_app():
    # Preserve keys for company info and proposal_settings.
    keys_to_preserve = {"client_info", "your_info", "proposal_settings"}
    for key in list(st.session_state.keys()):
        if key not in keys_to_preserve:
            del st.session_state[key]
    if "proposal_settings" not in st.session_state:
        st.session_state["proposal_settings"] = {"type": "Government", "tone": "Formal"}
    st.session_state["current_step"] = 1  # return to step 1
    save_state()
    try:
        st.rerun()
    except Exception:
        pass

def add_markdown_content(doc, md_text):
    """Add markdown-formatted text to a docx Document."""
    lines = md_text.splitlines()
    for line in lines:
        # Check for markdown headings (e.g., ### Heading)
        header_match = re.match(r'^(#{1,6})\s+(.*)', line)
        if header_match:
            level = len(header_match.group(1))
            text = header_match.group(2)
            doc.add_heading(text, level=level)
        else:
            # Replace bold syntax (**text**) with plain text 
            line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
            para = doc.add_paragraph(line)
            if para.runs:  # Check if there's at least one run
                para.runs[0].font.size = Pt(11)

# =============================================================================
# 1. Load Environment Variables & Initialize OpenAI
# =============================================================================
load_dotenv()

# Set up the Streamlit page configuration
st.set_page_config(page_title="Multi-Agent AI Proposal Drafting System", layout="wide")
load_state()

# Retrieve OpenAI API key from .env or prompt the user
openai_api_key = os.getenv("OPENAI_API_KEY")
if not openai_api_key:
    openai_api_key = st.text_input("Enter OpenAI API Key", type="password")
    if not openai_api_key:
        st.error("❌ OpenAI API Key is missing. Please enter it above.")
        st.stop()

# Initialize the coordinator and document extractor
coordinator = CoordinatorAgent(openai_api_key)
doc_extractor = DocumentExtractionAgent()

# =============================================================================
# 2. Sidebar Navigation & Reset Controls
# =============================================================================
st.sidebar.title("Proposal Navigation")
steps = [
    "1: Document Upload & Company Info",
    "2: Proposal Type & Structure",
    "3: Proposal Generation & Review",
    "4: Export & Handoff"
]

# Sidebar Start Over button
start_over_button = st.sidebar.button("Start Over", key="start_over")
if start_over_button:
    reset_app()

if "current_step" not in st.session_state:
    st.session_state["current_step"] = 1

selected_step_label = st.sidebar.radio("Go to Step", steps, index=st.session_state["current_step"] - 1)
selected_step = int(selected_step_label.split(":")[0].strip())
if selected_step != st.session_state["current_step"]:
    st.session_state["current_step"] = selected_step
    save_state()
    try:
        st.rerun()
    except AttributeError:
        pass

def next_step():
    if st.session_state["current_step"] < 5:
        st.session_state["current_step"] += 1
        save_state()

def prev_step():
    if st.session_state["current_step"] > 1:
        st.session_state["current_step"] -= 1
        save_state()
        try:
            st.rerun()
        except AttributeError:
            pass

# =============================================================================
# Section 1: Document Upload & Company Info
# =============================================================================
def section_1():
    st.title("Step 1: Document Upload & Analysis")
    
    # Initialize sources if not present
    if "sources" not in st.session_state:
        st.session_state["sources"] = {}
    
    col_left, col_right = st.columns(2)
            
    # LEFT PANEL: Client Organization Info
    with col_left:
        st.header("Client Organization Info")
        
        if "client_info" not in st.session_state:
            st.session_state["client_info"] = {"name": "", "description": "", "files": []}
        client = st.session_state["client_info"]
        client["name"] = st.text_input("Client Organization Name", value=client.get("name", ""))
        client["description"] = st.text_area("Client Organization Description", value=client.get("description", ""))
        
        uploaded_client_docs = st.file_uploader(
            "Upload Client Documents (RFP, SOW, etc.)",
            type=["pdf", "docx", "txt"],
            key="client_files",
            accept_multiple_files=True
        )
        
        client_files = []
        if uploaded_client_docs:
            st.markdown("### Uploaded Client Documents")
            for file in uploaded_client_docs:
                cols = st.columns([3, 3])
                with cols[0]:
                    st.write(file.name)
                with cols[1]:
                    category = st.selectbox(
                        "Category",
                        options=["RFP", "SOW", "Company Info", "Case Study", "News", "Whitepaper", "Other"],
                        key=f"client_category_{file.name}"
                    )
                text = doc_extractor.extract_text(file)
                if text:
                    client_files.append({"name": file.name, "text": text, "category": category})
                    key = "client_" + category
                    if key in st.session_state["sources"]:
                        st.session_state["sources"][key] += "\n" + text
                    else:
                        st.session_state["sources"][key] = text
                        
                    if "extracted_text" not in st.session_state:
                        st.session_state["extracted_text"] = text
                    else:
                        st.session_state["extracted_text"] += "\n\n" + text
            client["files"] = client_files
        st.session_state["client_info"] = client
    
    # RIGHT PANEL: Your Organization Info
    with col_right:
        st.header("Your Organization Info")
        if "your_info" not in st.session_state:
            st.session_state["your_info"] = {"name": "", "description": "", "files": []}
        your = st.session_state["your_info"]
        your["name"] = st.text_input("Your Organization Name", value=your.get("name", ""))
        your["description"] = st.text_area("Your Organization Description", value=your.get("description", ""))
        
        uploaded_your_docs = st.file_uploader(
            "Upload Your Organization Documents",
            type=["pdf", "docx", "txt"],
            key="your_files",
            accept_multiple_files=True
        )
        
        your_files = []
        if uploaded_your_docs:
            st.markdown("### Uploaded Company Documents")
            for file in uploaded_your_docs:
                cols = st.columns([3, 3])
                with cols[0]:
                    st.write(file.name)
                with cols[1]:
                    category = st.selectbox(
                        "Category",
                        options=["RFP", "SOW", "Company Info", "Case Study", "News", "Whitepaper", "Other"],
                        key=f"your_category_{file.name}"
                    )
                text = doc_extractor.extract_text(file)
                if text:
                    your_files.append({"name": file.name, "text": text, "category": category})
                    key = "company_" + category
                    if key in st.session_state["sources"]:
                        st.session_state["sources"][key] += "\n" + text
                    else:
                        st.session_state["sources"][key] = text
            your["files"] = your_files
        st.session_state["your_info"] = your

    # Navigation
    cols = st.columns([1, 1])
    with cols[0]:
        back_button = st.button("Back", key="nav_back_section1")
        if back_button:
            prev_step()
    with cols[1]:
        next_button = st.button("Next", key="nav_next_section1")
        if next_button:
            next_step()
            st.rerun()

# =============================================================================
# Section 2: Proposal Type & Structure
# =============================================================================
def section_2():
    st.title("Step 2: Proposal Type & Structure")

    # Navigation Buttons at Top Right
    nav_cols = st.columns(2)
    with nav_cols[0]:
        back_button = st.button("Go Back", key="nav_back_section2")
        if back_button:
            prev_step()
    with nav_cols[1]:
        next_button = st.button("Save & Next", key="nav_next_section2")
        if next_button:
            next_step()
    
    # Create two columns for better layout
    col_left, col_right = st.columns([1, 1])

    # LEFT PANEL: Set Proposal Type & Document Analysis
    with col_left:
        st.header("Set Proposal Type")
        if "proposal_settings" not in st.session_state:
            st.session_state["proposal_settings"] = {"type": "Government", "tone": "Formal"}
        settings = st.session_state["proposal_settings"]
        settings["type"] = st.radio("Proposal Type",
                                    ["Government", "Commercial", "Custom"],
                                    index=["Government", "Commercial", "Custom"].index(settings.get("type", "Government")),
                                    key="sec2_proposal_type")
        settings["tone"] = st.selectbox("Proposal Tone",
                                        ["Formal", "Persuasive", "Concise", "Custom"],
                                        index=["Formal", "Persuasive", "Concise", "Custom"].index(settings.get("tone", "Formal")),
                                        key="sec2_proposal_tone")

        # Document Analysis Section
        st.markdown("---")
        st.subheader("Document Analysis")
        
        # Analyze the documents
        extract_button = st.button("Extract Key Information", key="extract_key_info")
        has_analysis = "document_analysis" in st.session_state
        
        if extract_button or has_analysis:
            with st.spinner("Analyzing documents..."):
                if not has_analysis:
                    # Extract key requirements
                    requirements_result = coordinator.process_request(
                        "generate_technical_section",
                        client_text="Extract key requirements",
                        extracted_text=st.session_state["extracted_text"],
                        project_requirements="List the top 5-7 critical requirements from these documents. Focus on what the user needs to provide or clarify."
                    )
                    
                    # Store analysis results
                    st.session_state["document_analysis"] = {
                        "requirements": requirements_result.get("technical_solution", "")
                    }
        
        # Display the analysis results in a concise format
        if "document_analysis" in st.session_state:
            # Key Requirements
            st.markdown("### Critical Requirements & Missing Information")
            requirements = st.session_state["document_analysis"]["requirements"]
            # Extract bullet points from requirements
            bullet_points = [line.strip("- ").strip() for line in requirements.splitlines() if line.strip().startswith("-")]
            if bullet_points:
                for req in bullet_points[:5]:  # Show top 5 requirements
                    st.markdown(f"- {req}")
            else:
                st.info("No specific requirements found in the documents.")
        else:
            st.info("Click 'Extract Key Information' to analyze the documents.")

    # RIGHT PANEL: Proposal Structure
    with col_right:
        st.header("Proposal Structure")
        
        # Track if we have required data to generate a structure
        has_documents = "extracted_text" in st.session_state and st.session_state.get("extracted_text", "")
        has_structure = "proposal_sections" in st.session_state and len(st.session_state.get("proposal_sections", "").strip()) > 50
        
        # Status message about structure generation
        if not has_documents:
            st.warning("Please upload and process documents in Step 1 before generating a proposal structure.")
        elif not has_structure:
            st.info("Click 'Generate Proposal Structure' to create a customized outline based on your documents.")
        
        # Generate button for proposal structure - use a unique key not referenced in session state
        generate_structure = st.button("Generate Proposal Structure", key="generate_structure_button")
        
        if generate_structure and has_documents:
            with st.spinner("Analyzing documents and generating proposal structure..."):
                # Get client info and extracted text
                client_info = st.session_state["client_info"].get("description", "")
                extracted_text = st.session_state.get("extracted_text", "")
                
                # Step 1: First extract structured requirements from the RFP documents
                st.write("Step 1/2: Extracting key requirements from documents...")
                rfp_requirements = []
                
                # Check if we already extracted requirements
                if "extracted_requirements" in st.session_state and st.session_state["extracted_requirements"]:
                    rfp_requirements = st.session_state["extracted_requirements"]
                    st.write(f"Found {len(rfp_requirements)} previously extracted requirements.")
                else:
                    # Pre-process the text to identify key sections
                    # Look for common RFP section headers to focus extraction
                    headers = ["requirements", "scope of work", "deliverables", "specifications", 
                              "objectives", "timeline", "evaluation criteria", "qualifications"]
                    
                    # Find relevant sections in the text
                    important_sections = ""
                    lines = extracted_text.splitlines()
                    for i, line in enumerate(lines):
                        line_lower = line.lower()
                        # Check if line contains any of the headers
                        if any(header in line_lower for header in headers):
                            # Add this section (header + next 15 lines) to important sections
                            section_end = min(i + 15, len(lines))
                            important_sections += line + "\n" + "\n".join(lines[i+1:section_end]) + "\n\n"
                    
                    # If we found important sections, prioritize them
                    if important_sections:
                        extraction_text = important_sections + "\n\n" + extracted_text[:25000]
                    else:
                        extraction_text = extracted_text[:50000]
                    
                    # Extract requirements using a focused prompt
                    result = coordinator.process_request(
                        "generate_technical_section",
                        client_text="Requirements Extraction",
                        extracted_text=extraction_text,
                        project_requirements="""
                        You are an RFP specialist who needs to extract ALL specific requirements from an RFP document.
                        
                        Extract requirements in these categories:
                        1. TECHNICAL: What specific capabilities, features, and functionalities are required
                        2. PERFORMANCE: Required metrics, SLAs, speeds, capacities, or benchmarks
                        3. COMPLIANCE: Required standards, regulations, certifications, or policies
                        4. DELIVERABLES: Specific work products, documents, or artifacts required
                        5. TIMELINE: Deadlines, milestones, or schedule requirements
                        6. BUDGET: Cost constraints, payment structures, or financial terms
                        7. QUALIFICATIONS: Required experience, certifications, or staffing
                        
                        Format each requirement as a single-line bullet point starting with '-'.
                        Be concrete and specific. Do NOT include vague or generic statements.
                        Include ALL important requirements, even if there are many.
                        
                        Label each requirement with its category in [BRACKETS] at the start.
                        Example: 
                        - [TECHNICAL] System must support concurrent access by at least 500 users
                        - [TIMELINE] Final deliverables must be completed within 6 months of project start
                        """
                    )
                    
                    extracted_list = [line.strip("- ").strip() for line in 
                                    result.get("technical_solution", "").splitlines() 
                                    if line.strip().startswith("-")]
                    
                    if extracted_list:
                        rfp_requirements = extracted_list
                        st.session_state["extracted_requirements"] = rfp_requirements
                        st.write(f"Extracted {len(rfp_requirements)} key requirements from documents.")
                        
                        # Display the first few requirements for transparency
                        with st.expander("View extracted requirements"):
                            for req in rfp_requirements[:10]:
                                st.write(f"• {req}")
                            if len(rfp_requirements) > 10:
                                st.write(f"... and {len(rfp_requirements) - 10} more")
                    else:
                        st.write("No specific requirements found. Using general structure.")
                
                # Step 2: Generate a custom proposal structure based on the requirements
                st.write("Step 2/2: Creating targeted proposal structure based on requirements...")
                
                # Prepare requirements as formatted text for the prompt
                formatted_requirements = "\n".join([f"- {req}" for req in rfp_requirements])
                
                # Use the coordinator to generate a custom proposal structure
                result = coordinator.process_request(
                    "generate_technical_section",
                    client_text=client_info,
                    extracted_text=extracted_text[:10000],  # Use first 10k chars for context
                    project_requirements=f"""
                    As an expert RFP response strategist, create a proposal structure outline based on these SPECIFIC REQUIREMENTS:
                    
                    EXTRACTED RFP REQUIREMENTS:
                    {formatted_requirements}
                    
                    INSTRUCTIONS:
                    1. Create a logical structure with 5-7 main sections that collectively address ALL critical requirements
                    2. For each section, provide 3-5 specific bullet points that will be used to generate content
                    3. Each bullet point should clearly indicate which specific requirement(s) it addresses
                    4. Include specialized sections that match the project type evident in the requirements
                    5. Follow best practices for proposal structure:
                       - Start with an Executive Summary
                       - Include technical approach, implementation plan, qualifications, etc.
                       - End with a strong conclusion and next steps
                    
                    FORMAT:
                    - Section Title
                      * Bullet point 1 [Addresses requirement: specific requirement text]
                      * Bullet point 2 [Addresses requirements: specific requirement text]
                      * Bullet point 3 [Addresses requirement: specific requirement text]
                    
                    Ensure the entire structure forms a cohesive, logical flow and fully addresses the client's needs.
                    Use the client information to tailor the approach to their specific industry, size, and challenges.
                    """
                )
                
                # Extract the generated structure from the result
                generated_outline = result.get("technical_solution", "")
                
                # Debug information
                st.info(f"Generated outline length: {len(generated_outline)} characters")
                
                # If no structure was generated, use a basic template
                if not generated_outline or len(generated_outline.strip()) < 50:
                    generated_outline = """
- Executive Summary
  * Project goals and objectives
  * Value proposition
  * Key differentiators

- Technical Approach
  * Proposed solution overview
  * Key technologies
  * Integration approach

- Implementation Plan
  * Project phases
  * Timeline
  * Resource allocation

- Pricing
  * Cost breakdown
  * Payment schedule
  * ROI analysis

- Conclusion
  * Next steps
  * Contact information
"""
                
                # Update the session state with the generated structure
                st.session_state["proposal_sections"] = generated_outline
                # Add a timestamp to track when the structure was last modified
                st.session_state["structure_last_modified"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                # Flag to indicate a newly generated structure
                st.session_state["structure_newly_generated"] = True
                # Use rerun to update the UI immediately
                st.rerun()
        
        # Default outline only used if nothing exists in session state
        default_outline = """
- Executive Summary
  * Key project objectives and scope
  * Value proposition and benefits
  * Unique selling points
  * Project timeline overview

- Technical Approach
  * Solution architecture and components
  * Technology stack and tools
  * Integration points and APIs
  * Security and compliance measures
  * Scalability and performance considerations

- Pricing Proposal
  * Cost breakdown by phase
  * Resource allocation
  * Payment terms and conditions
  * ROI analysis and benefits

- Implementation Plan
  * Project phases and milestones
  * Team structure and roles
  * Risk mitigation strategies
  * Quality assurance process
  * Change management approach

- Compliance Statements
  * Regulatory requirements
  * Industry standards
  * Security certifications
  * Data protection measures

- Conclusion
  * Project success criteria
  * Long-term benefits
  * Next steps and call to action
"""
        # Display the outline in a text area for editing
        outline = st.text_area("Edit or add sections in your proposal outline:",
                              value=st.session_state.get("proposal_sections", default_outline),
                              height=400,
                              key="sec2_outline")
        st.session_state["proposal_sections"] = outline

        # Remove the duplicate text area - this whole section is replaced by the above code
        # Display the proposal structure text area
        if "proposal_sections" in st.session_state:
            # Display timestamp if available
            if "structure_last_modified" in st.session_state:
                st.caption(f"Last updated: {st.session_state['structure_last_modified']}")
            
            # Highlight the text area if newly generated
            if st.session_state.get("structure_newly_generated", False):
                st.success("✅ New proposal structure generated successfully!")
                # Reset the flag after displaying the message
                st.session_state["structure_newly_generated"] = False
        else:
            st.info("No proposal structure generated yet. Please generate a proposal structure in Step 2.")

# =============================================================================
# Section 3: Proposal Generation & Review
# =============================================================================
def section_3():
    st.title("Step 3: Proposal Generation & Review")
    
    nav_cols = st.columns(2)
    with nav_cols[0]:
        back_button = st.button("Go Back", key="nav_back_section3")
        if back_button:
            prev_step()
    with nav_cols[1]:
        next_button = st.button("Save & Next", key="nav_next_section3")
        if next_button:
            next_step()
            
    col_left, col_right = st.columns(2)
    
    # LEFT PANEL: Draft Proposal Content
    with col_left:
        st.header("Draft Proposal Content")
        
        # Get sections from the outline
        outline_text = st.session_state.get("proposal_sections", "")
        # Extract only main section titles (lines starting with - that don't have indentation)
        sections_list = []
        for line in outline_text.splitlines():
            if line.strip().startswith("-") and not line.startswith("  "):
                # Remove the dash and any leading/trailing whitespace
                section = line.lstrip("- ").strip()
                sections_list.append(section)
        
        # Initialize current section if needed
        if "current_section" not in st.session_state:
            st.session_state["current_section"] = 0
        
        # Section navigation
        section_nav = st.columns(2)
        with section_nav[0]:
            if st.button("Previous Section", key="prev_section_btn"):
                if st.session_state.get("current_section", 0) > 0:
                    st.session_state["current_section"] -= 1
        with section_nav[1]:
            if st.button("Next Section", key="next_section_btn"):
                if st.session_state.get("current_section", 0) < len(sections_list) - 1:
                    st.session_state["current_section"] += 1
                    
        # Section selection
        current_index = min(st.session_state.get("current_section", 0), len(sections_list)-1 if sections_list else 0)
        selected_section = st.radio("Select a section to edit:", sections_list, index=current_index) if sections_list else "No sections defined"
        if sections_list:
            st.session_state["current_section"] = sections_list.index(selected_section)
            current_section = selected_section
            
            # Get bullet points for the selected section
            section_bullets = []
            capture_bullets = False
            for line in outline_text.splitlines():
                line = line.strip()
                # Check if this is the main section we're looking for
                if line.startswith("-") and not line.startswith("  "):
                    section_title = line.lstrip("- ").strip()
                    # If we found our section, start capturing bullets
                    if section_title == current_section:
                        capture_bullets = True
                    # If we were capturing bullets and found a new section, stop capturing
                    elif capture_bullets:
                        capture_bullets = False
                # If we're in capture mode and this is a bullet point (indented and starts with *)
                elif capture_bullets and (line.startswith("*") or (line.startswith("  ") and "*" in line)):
                    bullet = line.lstrip(" *").strip()
                    if bullet:  # Only add non-empty bullets
                        section_bullets.append(bullet)
            
            # If no bullets were found, check if there are any indented lines that might be bullets
            if not section_bullets:
                capture_bullets = False
                for line in outline_text.splitlines():
                    line = line.strip()
                    # Check if this is the main section we're looking for
                    if line.startswith("-") and not line.startswith("  "):
                        section_title = line.lstrip("- ").strip()
                        # If we found our section, start capturing bullets
                        if section_title == current_section:
                            capture_bullets = True
                        # If we were capturing bullets and found a new section, stop capturing
                        elif capture_bullets:
                            capture_bullets = False
                    # If we're in capture mode and this is an indented line
                    elif capture_bullets and line.startswith("  "):
                        bullet = line.strip()
                        if bullet:  # Only add non-empty bullets
                            section_bullets.append(bullet)
            
            # Initialize generated sections dictionary if needed
            if "generated_sections" not in st.session_state:
                st.session_state["generated_sections"] = {}
                
            # Section generation buttons
            generation_options = st.columns(2)
            with generation_options[0]:
                gen_button = st.button(f"Generate {current_section}", key=f"gen_btn_{current_section}")
                if gen_button:
                    with st.spinner(f"Generating {current_section}..."):
                        extracted_text = st.session_state.get("extracted_text", "")
                        client_info = st.session_state["client_info"].get("description", "")
                        
                        # Create bullet points prompt if we have bullets
                        bullet_points_prompt = ""
                        if section_bullets:
                            bullet_points_prompt = "\n".join([f"- {bullet}" for bullet in section_bullets])
                        else:
                            bullet_points_prompt = f"- Content for {current_section}"
                        
                        # Different generation logic based on section type
                        if "Technical" in current_section:
                            result = coordinator.process_request(
                                "generate_technical_section",
                                client_text=client_info,
                                extracted_text=extracted_text,
                                project_requirements=f"Generate content ONLY for the '{current_section}' section. Focus exclusively on addressing these specific aspects without any introduction, summary, or conclusion:\n{bullet_points_prompt}\nEnsure the content directly addresses the requirements from the RFP and aligns with the user's organization capabilities."
                            )
                            st.session_state["generated_sections"][current_section] = result.get("technical_solution", "")
                        
                        elif "Timeline" in current_section or "Implementation" in current_section:
                            # Use bullet points to guide implementation plan
                            bullet_points_prompt = "\n".join([f"- {bullet}" for bullet in section_bullets])
                            result = coordinator.process_request(
                                "generate_timeline",
                                client_text=client_info,
                                relevant_text=extracted_text,
                                additional_context=f"Generate content ONLY for the '{current_section}' section. Focus exclusively on addressing these specific aspects without any introduction, summary, or conclusion:\n{bullet_points_prompt}\nEnsure the content directly addresses the requirements from the RFP and aligns with the user's organization capabilities."
                            )
                            st.session_state["generated_sections"][current_section] = result.get("timeline", "")
                        
                        elif "Pricing" in current_section:
                            # Use bullet points to guide pricing proposal
                            bullet_points_prompt = "\n".join([f"- {bullet}" for bullet in section_bullets])
                            result = coordinator.process_request(
                                "generate_technical_section",
                                client_text=client_info,
                                extracted_text=extracted_text,
                                project_requirements=f"Generate content ONLY for the '{current_section}' section. Focus exclusively on addressing these specific aspects without any introduction, summary, or conclusion:\n{bullet_points_prompt}\nEnsure the content directly addresses the requirements from the RFP and aligns with the user's organization capabilities."
                            )
                            st.session_state["generated_sections"][current_section] = result.get("technical_solution", "")
                        
                        elif "Executive Summary" in current_section:
                            # Special handling for Executive Summary
                            bullet_points_prompt = "\n".join([f"- {bullet}" for bullet in section_bullets])
                            result = coordinator.process_request(
                                "generate_technical_section",
                                client_text=client_info,
                                extracted_text=extracted_text,
                                project_requirements=f"Generate a concise Executive Summary that ONLY addresses these specific points without any additional introduction or conclusion:\n{bullet_points_prompt}\nKeep it focused on the client's needs and your organization's value proposition."
                            )
                            st.session_state["generated_sections"][current_section] = result.get("technical_solution", "")
                        
                        elif "Compliance" in current_section:
                            # Special handling for Compliance Statements
                            bullet_points_prompt = "\n".join([f"- {bullet}" for bullet in section_bullets])
                            result = coordinator.process_request(
                                "generate_technical_section",
                                client_text=client_info,
                                extracted_text=extracted_text,
                                project_requirements=f"Generate content ONLY for the '{current_section}' section that addresses these compliance aspects without any introduction or conclusion:\n{bullet_points_prompt}\nFocus on specific compliance statements relevant to the client's industry and requirements."
                            )
                            st.session_state["generated_sections"][current_section] = result.get("technical_solution", "")
                        
                        elif "Conclusion" in current_section:
                            # Special handling for Conclusion
                            bullet_points_prompt = "\n".join([f"- {bullet}" for bullet in section_bullets])
                            result = coordinator.process_request(
                                "generate_technical_section",
                                client_text=client_info,
                                extracted_text=extracted_text,
                                project_requirements=f"Generate content ONLY for the '{current_section}' section that addresses these specific points:\n{bullet_points_prompt}\nKeep it focused on next steps and call to action without summarizing the entire proposal."
                            )
                            st.session_state["generated_sections"][current_section] = result.get("technical_solution", "")
                        
                        else:
                            # Generic section generation using bullet points as guidance
                            bullet_points_prompt = "\n".join([f"- {bullet}" for bullet in section_bullets])
                            result = coordinator.process_request(
                                "generate_technical_section",
                                client_text=client_info,
                                extracted_text=extracted_text,
                                project_requirements=f"Generate content ONLY for the '{current_section}' section. Focus exclusively on addressing these specific aspects without any introduction, summary, or conclusion:\n{bullet_points_prompt}\nEnsure the content directly addresses the requirements from the RFP and aligns with the user's organization capabilities."
                            )
                            st.session_state["generated_sections"][current_section] = result.get("technical_solution", "")
            
            with generation_options[1]:
                clear_button = st.button("Clear Section", key=f"clear_btn_{current_section}")
                if clear_button:
                    if current_section in st.session_state["generated_sections"]:
                        del st.session_state["generated_sections"][current_section]
            
            # Display section for editing
            if current_section in st.session_state["generated_sections"]:
                edited_draft = st.text_area("Edit Section Content:",
                                        value=st.session_state["generated_sections"][current_section],
                                        height=400,  # Changed from default to match proposal structure
                                        key=f"edit_{current_section}")
                st.session_state["generated_sections"][current_section] = edited_draft
            else:
                st.info(f"Click 'Generate {current_section}' to create content for this section.")
                
            # User feedback section
            st.write("User Feedback:")
            user_feedback = st.text_area("Enter your feedback to improve the section content:", value="", height=100, key=f"feedback_{current_section}")
            feedback_button = st.button("Apply User Feedback", key=f"feedback_btn_{current_section}")
            if feedback_button:
                if current_section in st.session_state["generated_sections"]:
                    with st.spinner("Applying feedback..."):
                        current_draft = st.session_state["generated_sections"][current_section]
                        result = coordinator.process_request(
                            "apply_user_feedback",
                            current_draft=current_draft,
                            user_feedback=user_feedback
                        )
                        st.session_state["generated_sections"][current_section] = result.get("revised_draft", current_draft)
                        st.success("Feedback applied successfully!")
                else:
                    st.error("Please generate section content first before applying feedback.")
        else:
            st.error("No sections defined. Please define sections in Step 2.")
    
    with col_right:
        st.header("Live Proposal Preview")
        refresh_button = st.button("Refresh Preview", key="refresh_preview_btn")
        if refresh_button:
            pass  # Just triggers a rerun
            
        preview_container = st.empty()
        preview_content = ""
        for sec, content in st.session_state.get("generated_sections", {}).items():
            preview_content += f"## {sec}\n{content}\n\n"
            
        if not preview_content:
            preview_content = "No proposal content generated yet. Use the controls on the left to generate content for each section."
            
        preview_html = f'<div style="height:800px; overflow:auto; border:1px solid #ddd; padding:10px; white-space: pre-wrap;">{preview_content}</div>'
        st.markdown(preview_html, unsafe_allow_html=True)

# =============================================================================
# Section 4: Export & Handoff
# =============================================================================
def section_4():
    st.title("Step 4: Export & Handoff")
    
    # Navigation Button at Top Right
    nav_cols = st.columns([4, 1])
    with nav_cols[1]:
        back_button = st.button("Go Back", key="nav_back_section4")
        if back_button:
            prev_step()
            
    col_left, col_right = st.columns(2)
    
    # LEFT PANEL: Export Settings & Download
    with col_left:
        st.header("Export Settings & Final Touches")
        
        cover_page = st.text_input("Cover Page Text (optional):", value=st.session_state.get("cover_page", ""))
        st.session_state["cover_page"] = cover_page
        
        metadata = st.text_area("Enter metadata (submission date, contact info, version):",
                                value=st.session_state.get("metadata", ""))
        st.session_state["metadata"] = metadata

        # Generate the final proposal document for download
        gen_doc_button = st.button("Generate Final Document", key="gen_doc_btn")
        if gen_doc_button:
            with st.spinner("Generating document..."):
                final_doc = Document()
                
                # Add cover page information
                client = st.session_state.get("client_info", {})
                your = st.session_state.get("your_info", {})
                final_doc.add_paragraph(f"Prepared for: {client.get('name', 'Client')}")
                final_doc.add_paragraph(f"Prepared by: {your.get('name', 'Your Company')}")
                
                if st.session_state.get("cover_page"):
                    final_doc.add_paragraph(st.session_state["cover_page"])
                if st.session_state.get("metadata"):
                    final_doc.add_paragraph(st.session_state["metadata"])
                
                # Add each section
                for sec, content in st.session_state.get("generated_sections", {}).items():
                    final_doc.add_heading(sec, level=2)
                    add_markdown_content(final_doc, content)
                
                # Save document to a temporary file
                temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
                final_doc.save(temp_file.name)
                
                st.session_state["temp_docx"] = temp_file.name
                st.success("Document generated successfully!")
        
        # Download button
        if "temp_docx" in st.session_state:
            temp_docx_path = st.session_state["temp_docx"]
            try:
                # Check if the file exists
                if os.path.exists(temp_docx_path):
                    with open(temp_docx_path, "rb") as f:
                        st.download_button(
                            "Download Final Proposal (DOCX)",
                            f,
                            file_name="Final_Proposal.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="download_btn"
                        )
                else:
                    st.error("Document file not found. Please regenerate the document.")
                    # Remove the invalid reference
                    del st.session_state["temp_docx"]
            except Exception as e:
                st.error(f"Error accessing document file: {e}")
                # Remove the invalid reference
                if "temp_docx" in st.session_state:
                    del st.session_state["temp_docx"]
    
    # RIGHT PANEL: Required Documents Checklist
    with col_right:
        st.header("Required Documents Checklist")
        
        # Extract requirements if we haven't already
        if "extracted_requirements" not in st.session_state and "extracted_text" in st.session_state:
            with st.spinner("Extracting requirements..."):
                # Simulate requirement extraction for the checklist
                # This would typically use an agent to extract proper requirements
                extracted_text = st.session_state["extracted_text"]
                # Generate requirements list
                if coordinator and "openai_api_key" in locals() and openai_api_key:
                    result = coordinator.process_request(
                        "generate_technical_section",
                        client_text="Extract requirements for checklist",
                        extracted_text=extracted_text[:50000],  # Limit text length
                        project_requirements="List key requirements as bullet points"
                    )
                    tech_solution = result.get("technical_solution", "")
                    # Extract bullet points
                    extracted_list = [line.strip("- ").strip() for line in tech_solution.splitlines() 
                                     if line.strip().startswith("-")]
                    if extracted_list:
                        st.session_state["extracted_requirements"] = extracted_list
                    else:
                        st.session_state["extracted_requirements"] = [
                            "Executive Summary", 
                            "Technical Approach", 
                            "Pricing Details",
                            "Implementation Timeline",
                            "Team Qualifications",
                            "References"
                        ]
        
        # Display checklist
        if "extracted_requirements" in st.session_state:
            st.write("Check off required elements that have been completed:")
            for i, req in enumerate(st.session_state["extracted_requirements"]):
                st.checkbox(req, value=False, key=f"check_{i}")
        else:
            st.info("No requirements extracted. Complete previous steps to generate the checklist.")

# =============================================================================
# Main App Dispatcher
# =============================================================================
def main():
    st.title("Multi-Agent AI Proposal Drafting System")
    current = st.session_state.get("current_step", 1)
    if current == 1:
        section_1()
    elif current == 2:
        section_2()
    elif current == 3:
        section_3()
    elif current == 4:
        section_4()
    save_state()

if __name__ == "__main__":
    main()
